#!/usr/bin/env python3
"""Scrape SAT Reading & Writing knowledge-point practice questions.

Sources:
  - 历年真题
  - Educator Question Bank

Outputs a Word document with stem, options, answer, text analysis,
category path, and difficulty for each question.
"""

from __future__ import annotations

import argparse
import html
import json
import re
import sys
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import requests

try:
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
except ImportError:  # pragma: no cover
    Document = None  # type: ignore

BASE_URL = "https://www.kaoshixiansheng.top/server"
SUBJECT = "ReadingWriting"
TARGET_BANKS = ["历年真题", "Educator Question Bank"]
DIFFICULTIES = ["简单", "中等", "复杂"]

DEFAULT_OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "research"
    / "data"
    / "sat_rw_knowledge_questions.docx"
)
CHECKPOINT_PATH = (
    Path(__file__).resolve().parents[1]
    / "research"
    / "data"
    / "sat_rw_knowledge_questions.json"
)


@dataclass
class QuestionItem:
    bank_type: str
    level1: str
    level2: str
    level3: str
    knowledge_id: int
    knowledge_name: str
    difficulty: str
    knowledge_group: int
    knowledge_group_name: str
    question_id: int
    question_number: int
    question_type: str
    article: str
    stem: str
    options: list[dict[str, str]]
    answer: str = ""
    text_analysis: str = ""
    record_id: Any = None


@dataclass
class GroupKey:
    bank_type: str
    knowledge_id: int
    knowledge_name: str
    level1: str
    level2: str
    level3: str
    difficulty: str
    knowledge_group: int
    knowledge_group_name: str
    exercise_num: int = 0


def login(session: requests.Session, identity: str, password: str) -> str:
    response = session.post(
        f"{BASE_URL}/user/login",
        params={
            "type": 1,
            "rememberme": "false",
            "identity": identity,
            "principal": password,
        },
        headers={"Content-Type": "application/x-www-form-urlencoded;charset=utf-8"},
        timeout=30,
    )
    response.raise_for_status()
    payload = response.json()
    if payload.get("status") != "success":
        raise RuntimeError(f"Login failed: {payload}")
    token = response.headers.get("Authorization")
    if not token:
        raise RuntimeError("Login succeeded but no Authorization header returned")
    return token


def api_get(session: requests.Session, path: str, params: dict | None = None) -> Any:
    response = session.get(
        f"{BASE_URL}{path}",
        params=params or {},
        timeout=60,
    )
    response.raise_for_status()
    payload = response.json()
    if payload.get("status") != "success":
        raise RuntimeError(f"GET {path} failed: {payload}")
    return payload.get("data")


def api_post(session: requests.Session, path: str, data: dict | None = None) -> Any:
    response = session.post(
        f"{BASE_URL}{path}",
        json=data or {},
        headers={"Content-Type": "application/json;charset=UTF-8"},
        timeout=60,
    )
    response.raise_for_status()
    payload = response.json()
    if payload.get("status") != "success":
        raise RuntimeError(f"POST {path} failed: {payload}")
    return payload.get("data")


def strip_html(raw: str | None) -> str:
    if not raw:
        return ""
    text = raw
    text = re.sub(r"(?i)<\s*br\s*/?\s*>", "\n", text)
    text = re.sub(r"(?i)</\s*p\s*>", "\n", text)
    text = re.sub(r"(?i)</\s*div\s*>", "\n", text)
    text = re.sub(r"(?i)</\s*li\s*>", "\n", text)
    text = re.sub(r"<[^>]+>", "", text)
    text = html.unescape(text)
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def walk_groups(nodes: list[dict], path: list[str] | None = None) -> list[GroupKey]:
    path = path or []
    groups: list[GroupKey] = []
    for node in nodes or []:
        name = node.get("name") or ""
        cur = path + [name]
        children = node.get("children") or []
        if children:
            groups.extend(walk_groups(children, cur))
            continue

        level1 = cur[0] if len(cur) >= 1 else ""
        level2 = cur[1] if len(cur) >= 2 else ""
        level3 = cur[2] if len(cur) >= 3 else ""
        knowledge_name = level3 or level2 or level1
        knowledge_id = int(node.get("id"))

        for exercise in node.get("exercises") or []:
            difficulty = exercise.get("difficulty") or ""
            for g in exercise.get("groups") or []:
                groups.append(
                    GroupKey(
                        bank_type="",  # filled by caller
                        knowledge_id=int(g.get("knowledgeId") or knowledge_id),
                        knowledge_name=knowledge_name,
                        level1=level1,
                        level2=level2,
                        level3=level3,
                        difficulty=difficulty,
                        knowledge_group=int(g.get("knowledgeGroup")),
                        knowledge_group_name=g.get("knowledgeGroupName") or "",
                        exercise_num=int(g.get("exerciseNum") or 0),
                    )
                )
    return groups


def collect_catalog(session: requests.Session, banks: list[str]) -> list[GroupKey]:
    catalog: list[GroupKey] = []
    for bank in banks:
        tree = api_get(
            session,
            "/sat/student/paper/exercise/base",
            {
                "subject": SUBJECT,
                "bankType": bank,
                "knowledgeId": "",
            },
        )
        groups = walk_groups(tree if isinstance(tree, list) else [])
        for g in groups:
            g.bank_type = bank
        catalog.extend(groups)
        print(f"Catalog {bank}: {len(groups)} practice groups")
    return catalog


def fetch_group_questions(
    session: requests.Session, group: GroupKey
) -> list[dict]:
    data = api_get(
        session,
        "/sat/student/paper/exercise/get",
        {
            "bankType": group.bank_type,
            "subject": SUBJECT,
            "knowledgeId": group.knowledge_id,
            "knowledgeGroup": group.knowledge_group,
            "difficulty": group.difficulty,
        },
    )
    return data or []


def create_serial(session: requests.Session, group: GroupKey) -> str:
    data = api_post(
        session,
        "/sat/student/paper/serialnumberid",
        {
            "paperId": group.knowledge_group,
            "name": group.knowledge_group_name or f"练习{group.knowledge_group}",
        },
    )
    return str(data)


def commit_group(
    session: requests.Session,
    group: GroupKey,
    serial_number_id: str,
    questions: list[dict],
) -> None:
    payload = {
        "paperId": group.knowledge_group,
        "serialNumberId": serial_number_id,
        "bankType": group.bank_type,
        "practiceScene": "KNOWLEDGE",
        "commitBaseDTOList": [
            {
                "module": q.get("module") or "M1",
                "questionId": q.get("questionId"),
                "selfChoiceAnswer": "",
                "selfFillInAnswer": "",
                "usedTime": 0,
                "tag": False,
                "everMarked": False,
            }
            for q in questions
        ],
    }
    api_post(session, "/sat/student/paper/exercise/commit", payload)


def fetch_statistics(session: requests.Session, serial_number_id: str) -> dict:
    data = api_get(
        session,
        "/sat/student/paper/exercise/record/statistics",
        {"serialNumberId": serial_number_id},
    )
    return data or {}


def fetch_record_detail(session: requests.Session, record_id: Any) -> dict:
    data = api_get(
        session,
        "/sat/student/paper/question/getbyrecordId",
        {"recordId": record_id},
    )
    return data or {}


def extract_question_details(stats: dict) -> list[dict]:
    """Return ordered questionDetails rows from exercise statistics."""
    if not stats:
        return []
    details = stats.get("questionDetails")
    if isinstance(details, list):
        return [row for row in details if isinstance(row, dict)]
    return []


def enrich_with_answers(
    session: requests.Session,
    group: GroupKey,
    questions: list[dict],
    sleep_s: float,
) -> list[QuestionItem]:
    items: list[QuestionItem] = []
    for q in questions:
        options = q.get("options") or []
        norm_options = []
        for opt in options:
            if isinstance(opt, dict):
                norm_options.append(
                    {
                        "mark": str(opt.get("mark") or ""),
                        "content": strip_html(opt.get("content") or ""),
                    }
                )
            else:
                norm_options.append({"mark": "", "content": strip_html(str(opt))})

        items.append(
            QuestionItem(
                bank_type=group.bank_type,
                level1=group.level1,
                level2=group.level2,
                level3=group.level3,
                knowledge_id=group.knowledge_id,
                knowledge_name=group.knowledge_name,
                difficulty=group.difficulty,
                knowledge_group=group.knowledge_group,
                knowledge_group_name=group.knowledge_group_name,
                question_id=int(q.get("questionId")),
                question_number=int(q.get("questionNumber") or 0),
                question_type=str(q.get("type") or ""),
                article=strip_html(q.get("article") or ""),
                stem=strip_html(q.get("stem") or ""),
                options=norm_options,
            )
        )

    # Create a practice session and commit blank answers to unlock keys/analysis.
    serial_id = create_serial(session, group)
    time.sleep(sleep_s)
    commit_group(session, group, serial_id, questions)
    time.sleep(sleep_s)
    stats = fetch_statistics(session, serial_id)
    details = extract_question_details(stats)

    # Prefer matching by questionNumber; otherwise zip by stable order.
    by_number: dict[int, dict] = {}
    for row in details:
        num = row.get("questionNumber")
        if num is not None:
            by_number[int(num)] = row

    for idx, item in enumerate(items):
        row = by_number.get(item.question_number)
        if row is None and idx < len(details):
            row = details[idx]
        row = row or {}

        item.answer = str(row.get("correctAnswer") or row.get("answer") or "")
        item.record_id = row.get("recordId") or row.get("id")

        if not item.record_id:
            continue

        try:
            detail = fetch_record_detail(session, item.record_id)
            time.sleep(sleep_s)
        except Exception as exc:  # noqa: BLE001
            print(f"  warn getbyrecordId {item.record_id}: {exc}")
            detail = {}

        if not detail:
            continue

        detail_qid = detail.get("questionId")
        if detail_qid is not None and int(detail_qid) != item.question_id:
            print(
                f"  warn qid mismatch record={item.record_id} "
                f"expected={item.question_id} got={detail_qid}"
            )

        item.text_analysis = strip_html(
            detail.get("textAnalysis") or detail.get("analysis") or ""
        )
        if not item.answer:
            item.answer = str(
                detail.get("correctAnswer")
                or detail.get("answer")
                or detail.get("rightAnswer")
                or ""
            )

    return items


def category_label(item: QuestionItem) -> str:
    parts = [p for p in [item.level1, item.level2, item.level3] if p]
    return " > ".join(parts) if parts else item.knowledge_name


def save_checkpoint(items: list[QuestionItem], path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    payload = [item.__dict__ for item in items]
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def load_checkpoint(path: Path) -> list[QuestionItem]:
    if not path.exists():
        return []
    raw = json.loads(path.read_text(encoding="utf-8"))
    items = []
    for row in raw:
        items.append(QuestionItem(**row))
    return items


def set_run_font(run, name: str = "Microsoft YaHei", size: int = 11) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def add_heading_line(doc: Document, text: str, size: int = 16, bold: bool = True) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size)
    run.bold = bold


def add_body(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    lab = p.add_run(f"{label}")
    set_run_font(lab, size=11)
    lab.bold = True
    if text:
        body = p.add_run(text if text.startswith("\n") else f"\n{text}")
        set_run_font(body, size=11)


def write_docx(items: list[QuestionItem], output: Path) -> None:
    if Document is None:
        raise RuntimeError("python-docx is required. Install with: pip install python-docx")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(11)

    add_heading_line(doc, "SAT 阅读与写作 · 知识点练习题库", size=18)
    p = doc.add_paragraph()
    run = p.add_run(
        f"题库范围：{'、'.join(TARGET_BANKS)}　共 {len(items)} 题\n"
        "字段：类别、难度、题干、选项、答案、文字解析"
    )
    set_run_font(run, size=11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    current_bank = None
    for idx, item in enumerate(items, start=1):
        if item.bank_type != current_bank:
            current_bank = item.bank_type
            doc.add_page_break()
            add_heading_line(doc, f"题库：{current_bank}", size=16)

        add_heading_line(doc, f"{idx}. [{item.bank_type}] 第{item.question_number}题", size=13)
        add_body(doc, "类别：", category_label(item))
        add_body(
            doc,
            "难度 / 练习组：",
            f"{item.difficulty}　{item.knowledge_group_name or f'练习{item.knowledge_group}'}",
        )
        if item.article:
            add_body(doc, "材料：", item.article)
        add_body(doc, "题干：", item.stem)
        if item.options:
            opt_lines = []
            for opt in item.options:
                mark = opt.get("mark") or ""
                content = opt.get("content") or ""
                opt_lines.append(f"{mark}. {content}" if mark else content)
            add_body(doc, "选项：", "\n".join(opt_lines))
        add_body(doc, "答案：", item.answer or "（未获取）")
        add_body(doc, "文字解析：", item.text_analysis or "（无文字解析）")
        doc.add_paragraph("")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def is_auth_error(exc: Exception) -> bool:
    text = str(exc)
    return "未登陆" in text or "100001" in text or "permission_err" in text


def scrape(
    identity: str,
    password: str,
    banks: list[str],
    output: Path,
    checkpoint: Path,
    sleep_s: float,
    limit_groups: int | None,
    resume: bool,
) -> list[QuestionItem]:
    session = requests.Session()
    session.headers.update(
        {
            "User-Agent": (
                "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/120.0.0.0 Safari/537.36"
            )
        }
    )

    def ensure_login() -> None:
        token = login(session, identity, password)
        session.headers["Authorization"] = token
        print("Login OK")

    ensure_login()

    catalog = collect_catalog(session, banks)
    if limit_groups is not None:
        catalog = catalog[:limit_groups]
        print(f"Limited to first {len(catalog)} groups")

    items: list[QuestionItem] = load_checkpoint(checkpoint) if resume else []
    done_keys = {
        (
            i.bank_type,
            i.knowledge_id,
            i.difficulty,
            i.knowledge_group,
            i.question_id,
        )
        for i in items
    }
    done_groups = {
        (i.bank_type, i.knowledge_id, i.difficulty, i.knowledge_group) for i in items
    }
    print(f"Resuming with {len(items)} questions, {len(done_groups)} groups done")

    for gi, group in enumerate(catalog, start=1):
        gkey = (
            group.bank_type,
            group.knowledge_id,
            group.difficulty,
            group.knowledge_group,
        )
        if gkey in done_groups:
            continue

        label = (
            f"[{gi}/{len(catalog)}] {group.bank_type} | "
            f"{group.level1}>{group.level2}>{group.level3} | "
            f"{group.difficulty} | {group.knowledge_group_name}"
        )
        print(label)
        try:
            questions = fetch_group_questions(session, group)
            time.sleep(sleep_s)
            if not questions:
                print("  empty group")
                done_groups.add(gkey)
                continue
            group_items = enrich_with_answers(session, group, questions, sleep_s)
            for item in group_items:
                key = (
                    item.bank_type,
                    item.knowledge_id,
                    item.difficulty,
                    item.knowledge_group,
                    item.question_id,
                )
                if key in done_keys:
                    continue
                items.append(item)
                done_keys.add(key)
            done_groups.add(gkey)
            save_checkpoint(items, checkpoint)
            print(
                f"  +{len(group_items)} questions "
                f"(answers={sum(1 for x in group_items if x.answer)}/"
                f"analysis={sum(1 for x in group_items if x.text_analysis)}) "
                f"total={len(items)}"
            )
        except Exception as exc:  # noqa: BLE001
            print(f"  ERROR: {exc}")
            save_checkpoint(items, checkpoint)
            if is_auth_error(exc):
                try:
                    ensure_login()
                except Exception as login_exc:  # noqa: BLE001
                    print(f"  re-login failed: {login_exc}")
                    break
            time.sleep(max(sleep_s, 1.0))

    write_docx(items, output)
    save_checkpoint(items, checkpoint)
    print(f"\nDone. {len(items)} questions")
    print(f"Word: {output}")
    print(f"JSON: {checkpoint}")
    return items


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("identity", nargs="?", default="18910210028")
    parser.add_argument("password", nargs="?", default="123456")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--checkpoint", type=Path, default=CHECKPOINT_PATH)
    parser.add_argument("--sleep", type=float, default=0.25)
    parser.add_argument("--limit-groups", type=int, default=None)
    parser.add_argument("--resume", action="store_true")
    parser.add_argument(
        "--banks",
        nargs="+",
        default=TARGET_BANKS,
        help="Bank types to scrape",
    )
    parser.add_argument(
        "--docx-only",
        action="store_true",
        help="Rebuild Word document from existing checkpoint JSON only",
    )
    args = parser.parse_args()

    if args.docx_only:
        items = load_checkpoint(args.checkpoint)
        if not items:
            print(f"No checkpoint at {args.checkpoint}")
            return 1
        write_docx(items, args.output)
        print(f"Wrote {len(items)} questions to {args.output}")
        return 0

    scrape(
        identity=args.identity,
        password=args.password,
        banks=args.banks,
        output=args.output,
        checkpoint=args.checkpoint,
        sleep_s=args.sleep,
        limit_groups=args.limit_groups,
        resume=args.resume,
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
